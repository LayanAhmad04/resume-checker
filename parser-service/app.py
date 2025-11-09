from flask import Flask, request, jsonify
import os, re, json, base64, tempfile
import fitz
import docx
import spacy
import psycopg2
from config import DB_DSN, OPENAI_KEY, PARSER_PORT
from openai import OpenAI

client = OpenAI(api_key=OPENAI_KEY)
nlp = spacy.load("en_core_web_sm")

app = Flask(__name__)

UPLOAD_DIR = os.path.abspath(
    os.getenv("UPLOAD_DIR", os.path.join(os.path.dirname(__file__), "..", "uploads"))
)
print("Parser service UPLOAD_DIR:", UPLOAD_DIR)


def extract_text_from_pdf(path):
    doc = fitz.open(path)
    return "\n".join([p.get_text() or "" for p in doc])


def extract_text_from_docx(path):
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text])


def extract_text(path):
    try:
        low = path.lower()
        if low.endswith(".pdf"):
            import fitz
            doc = fitz.open(path)
            return "\n".join([p.get_text() or "" for p in doc])
        elif low.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        elif low.endswith(".doc"):
            with open(path, "rb") as f:
                return f.read().decode("utf-8", errors="ignore")
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        print("extract_text error:", e)
        return ""


def clean_rtf(text):
    """Remove basic RTF control sequences and braces."""
    return re.sub(r"\\[a-z]+\d*|[{}]", "", text)


# name and email extraction
def extract_name_email(text, file_ext=None):
    # Extract email
    emails = re.findall(r"[\w\.-]+@[\w\.-]+", text)
    email = emails[0] if emails else None

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # --- Clean font/metadata words in DOC ---
    if file_ext and file_ext.lower().endswith(".doc"):
        font_keywords = r"\b(?:calibri|arial|times new roman|cambria|courier new|verdana|tahoma|georgia|helvetica)\b"
        lines = [re.sub(font_keywords, "", l, flags=re.I).strip() for l in lines]
        lines = [l for l in lines if l]

    candidate_name = None

    blacklist_keywords = r"\b(?:phone|email|linkedin|cv|resume|profile|skills|experience|projects|education|machine learning|python|docker|algorithms|professional summary|summary|objective|contact)\b"

    def clean_line(line):
        return re.sub(r"[^A-Za-z\s]", "", line).strip()

    for line in lines[:20]:
        if re.search(blacklist_keywords, line, re.I):
            continue
        cline = clean_line(line)
        words = cline.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words):
            candidate_name = cline.title()
            break

    if not candidate_name:
        clean_text = "\n".join(lines[:40])
        doc = nlp(clean_text)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                words = name.split()
                if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words):
                    if not re.search(blacklist_keywords, name, re.I):
                        candidate_name = name
                        break

    if not candidate_name and email:
        before_email = text.split(email)[0]
        match = re.search(r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})", before_email)
        if match and not re.search(blacklist_keywords, match.group(1), re.I):
            candidate_name = match.group(1)

    return candidate_name, email

# database connection
def db_connect():
    return psycopg2.connect(
        host=DB_DSN["host"],
        port=DB_DSN["port"],
        user=DB_DSN["user"],
        password=DB_DSN["password"],
        dbname=DB_DSN["database"],
    )


# openAI scoring logic
def call_openai_subscores_and_justification(job_description: str, resume_text: str, criteria: dict):
    weights = {}
    total_raw = 0.0
    for k, v in (criteria or {}).items():
        try:
            total_raw += float(v)
        except Exception:
            total_raw += 0.0

    if total_raw <= 0:
        n = len(criteria) or 1
        for k in (criteria or {}):
            weights[k] = 1.0 / n
    else:
        for k, v in (criteria or {}).items():
            weights[k] = float(v) / total_raw

    prompt = f"""
You are an objective HR scoring assistant. Given the job description and candidate resume, produce ONLY valid JSON matching this format:

{{
  "subscores": {{
    "<criterion>": {{ "score": 0.00, "reason": "short justification" }}
  }},
  "contributions": {{
    "<criterion>": 0.00
  }},
  "total_score_out_of_10": 0.00,
  "overall_justification": "1–3 sentences summarizing final score reasoning."
}}

Rules:
- Use each provided criterion name exactly.
- Score: 0.00–1.00 (two decimals)
- contribution = score * normalized_weight * 10
- total_score_out_of_10 = sum(contributions)
- Return valid JSON only.

Job description:
\"\"\"{job_description}\"\"\"

Resume (truncated):
\"\"\"{resume_text[:7000]}\"\"\"

Normalized weights:
{json.dumps(weights, indent=2)}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a concise, objective HR scoring assistant that outputs only valid JSON.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            max_tokens=700,
            seed=42
        )

        txt = response.choices[0].message.content
        first, last = txt.find("{"), txt.rfind("}")
        if first == -1 or last == -1 or last < first:
            raise ValueError("No JSON found in response")

        j = json.loads(txt[first:last + 1])
        subs = {}
        contribs = {}

        for k in weights.keys():
            entry = j.get("subscores", {}).get(k)
            s = float(entry["score"]) if entry and "score" in entry else 0.0
            reason = entry.get("reason", "") if entry else ""
            s = max(0.0, min(1.0, round(s, 2)))
            subs[k] = {"score": s, "reason": reason}
            contribs[k] = round(s * weights[k] * 10, 2)

        total = round(sum(contribs.values()), 2)
        overall = j.get("overall_justification", "")
        return subs, contribs, total, overall

    except Exception as e:
        print("OpenAI scoring/parsing error:", e)
        subs_fallback, contribs_fallback = {}, {}
        for k, w in (criteria or {}).items():
            subs_fallback[k] = {"score": 0.5, "reason": "Fallback neutral score"}
            contribs_fallback[k] = round(0.5 * (float(w) / (total_raw or 1)) * 10, 2)
        total_fallback = round(sum(contribs_fallback.values()), 2)
        return (
            subs_fallback,
            contribs_fallback,
            total_fallback,
            "Fallback justification: OpenAI failed or invalid response.",
        )


# main parser endpoint
@app.route("/process", methods=["POST"])
def process():
    data = request.json
    jobId = data.get("jobId")
    candidateId = data.get("candidateId")
    filename = data.get("filename")
    file_data = data.get("fileData")
    text_data = data.get("textData")

    if not (jobId and candidateId and filename):
        return jsonify({"error": "missing parameters"}), 400

    text = ""

    if text_data:
        text = text_data
        print(f"Processing candidate {candidateId} using raw_text input")
    elif file_data:
        try:
            temp_path = os.path.join(tempfile.gettempdir(), filename)
            with open(temp_path, "wb") as f:
                f.write(base64.b64decode(file_data))
            text = extract_text(temp_path)

            if filename.lower().endswith(".doc"):
                text = clean_rtf(text)

        except Exception as e:
            return jsonify({"error": "failed to decode file", "details": str(e)}), 400
        finally:
            try:
                os.remove(temp_path)
            except Exception:
                pass
    else:
        return jsonify({"error": "no fileData or textData provided"}), 400

    name, email = extract_name_email(text, filename)


    conn = db_connect()
    cur = conn.cursor()
    cur.execute("SELECT description, criteria FROM jobs WHERE id=%s", (jobId,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "job not found"}), 404
    job_description, criteria_json = row
    criteria = json.loads(criteria_json) if isinstance(criteria_json, str) else criteria_json or {}

    try:
        subscores, contributions, total_out_of_10, overall_just = call_openai_subscores_and_justification(
            job_description, text, criteria
        )
    except Exception as e:
        print("OpenAI helper failed:", e)
        subscores = {k: {"score": 0.5, "reason": "Fallback scoring"} for k in criteria.keys()}
        contributions = {k: 0.5 for k in criteria.keys()}
        total_out_of_10 = 5.0
        overall_just = "Fallback justification."

    score = total_out_of_10

    try:
        cur.execute(
            """
            UPDATE candidates
            SET name=%s,
                email=%s,
                raw_text=%s,
                parsed_data=%s,
                subscores=%s,
                score=%s,
                justification=%s,
                processed_at=NOW()
            WHERE id=%s
        """,
            (
                name,
                email,
                text[:10000],
                json.dumps({"name": name, "email": email}),
                json.dumps(subscores),
                score,
                json.dumps({"overall": overall_just, "contributions": contributions}),
                candidateId,
            ),
        )
        conn.commit()
    except Exception as e:
        print("DB update failed:", e)
        conn.rollback()
        conn.close()
        return jsonify({"error": "db update failed", "details": str(e)}), 500

    conn.close()
    return jsonify(
        {
            "ok": True,
            "candidateId": candidateId,
            "name": name,
            "email": email,
            "score": score,
            "subscores": subscores,
            "justification": overall_just,
            "contributions": contributions,
        }
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PARSER_PORT, debug=True)
